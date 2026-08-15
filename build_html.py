"""
build_html.py, revised 19 July 2026.

Renders CV and cover letter per role. CV layout matches the Ojas Indulkar Lebenslauf
reference Rah uploaded on 19 July 2026, per the 19 July CLAUDE.md rule that supersedes
the 18 July Lebenslauf rule.

Deliverables per role, under drafts/[folder]/:
- CV_Rahul_Rawat.html
- CV_Rahul_Rawat.pdf
- CV_Rahul_Rawat.docx
- CoverLetter_Rahul_Rawat.docx
- CoverLetter_Rahul_Rawat.md
- CV_Rahul_Rawat.md
"""

import os
import re
import sys
from pathlib import Path
from html import escape
from datetime import date
import weasyprint
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Mm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
DRAFTS = ROOT / "drafts"

# --- Palette --------------------------------------------------------------
NAVY_HEX = "#1F3A5F"
NAVY_GREY_HEX = "#6C7A93"
BODY_HEX = "#1F2937"
RULE_HEX = "#B7C0CE"
RUST_HEX = "#B8663F"

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
NAVY_GREY = RGBColor(0x6C, 0x7A, 0x93)
BODY = RGBColor(0x1F, 0x29, 0x37)
RUST = RGBColor(0xB8, 0x66, 0x3F)


# --- Personal Details values, 19 July 2026 rule ---------------------------
PD_FIELDS_EN = [
    ("Address", "C2 16, 68159 Mannheim, Germany"),
    ("Phone", "015563603340"),
    ("Email", "rahulrawat2r@gmail.com"),
    ("LinkedIn", "linkedin.com/in/rahulrawat2r"),
    ("GitHub", "github.com/rahulrawat20022002"),
    ("Portfolio", "rah-portfolio.pages.dev"),
    ("Date of birth", "20 February 2002"),
    ("Nationality", "Indian, student visa with valid work permit"),
    ("Availability", "Werkstudent 20 hours per week immediately, full time from April 2027"),
]

PD_FIELDS_DE = [
    ("Adresse", "C2 16, 68159 Mannheim, Deutschland"),
    ("Telefon", "015563603340"),
    ("E-Mail", "rahulrawat2r@gmail.com"),
    ("LinkedIn", "linkedin.com/in/rahulrawat2r"),
    ("GitHub", "github.com/rahulrawat20022002"),
    ("Portfolio", "rah-portfolio.pages.dev"),
    ("Geburtsdatum", "20 February 2002"),
    ("Nationalität", "Indisch, Studentenvisum mit gültiger Arbeitserlaubnis"),
    ("Verfügbarkeit", "Werkstudent 20 Stunden pro Woche sofort, Vollzeit ab April 2027"),
]

# --- Section heading labels ----------------------------------------------
HDR_EN = {
    "pd": "Personal Details",
    "profile": "Profile",
    "experience": "Professional Experience",
    "education": "Education",
    "projects": "Personal Projects",
    "research": "Research and Thesis",
    "certifications": "Certifications",
    "achievements": "Achievements",
    "languages": "Languages",
}
HDR_DE = {
    "pd": "Persönliche Daten",
    "profile": "Profil",
    "experience": "Berufserfahrung",
    "education": "Ausbildung",
    "projects": "Persönliche Projekte",
    "research": "Forschung und Abschlussarbeit",
    "certifications": "Zertifikate",
    "achievements": "Auszeichnungen",
    "languages": "Sprachen",
}


# --- Bold inline metrics --------------------------------------------------
# Match numeric quantities with meaning: percentages, ranges, multipliers,
# durations, magnitudes. Deliberately excludes bare years and reference IDs.
# Supports both English units and German units (Prozent, tausend, Monaten,
# Datensaetzen, Modelle, Faktor, fach, Jahre, Zertifikate, etc.) and both
# English decimal points (0.79) and German decimal commas (0,79).
_UNIT_WORDS = (
    r"percent|%|thousand|million|billion|K\b|M\b|B\b|"
    r"hours?|days?|weeks?|minutes?|seconds?|months?|years?|"
    r"times|fold|records?|companies|files?|classifiers?|"
    r"models?|patients?|folds?|euros?|gbps|mbps|gb|mb|kb|"
    r"chunks?|agents?|nodes?|jobs?|route[ns]?|dimensions?|documents?|dokumente?[nr]?|metrics?|metrike[nr]?|"
    # CV vocabulary that commonly follows a count
    r"tier[ns]?|table[ns]?|tabelle[nr]?|layer[ns]?|schicht[eanuers]*|"
    r"page[ns]?|seite[nr]?|indicator[ns]?|indikator(?:en)?|"
    r"target[ns]?|ziel[eanuers]*|feature[ns]?|candidate[ns]?|kandidate[nr]?|"
    r"pass(?:es)?|paesse|paess|quantile[ns]?|quantil(?:e[nr]?)?|"
    r"sensor(?:e[nr]?)?|sensors?|proxy|proxies|proxie[ns]?|"
    r"macro[ns]?|makros?|module?|module[nr]?|modul(?:e[nr]?)?|"
    r"class(?:es)?|klasse[nr]?|iteration[ns]?|iteratione[nr]?|"
    r"forecast[ns]?|prognose[nr]?|sprint[ns]?|"

    # German unit words
    r"Prozent|tausend|Millionen|Milliarden|Millionaerinnen|"
    r"Stunden?|Tage|Tagen|Wochen|Woche|Minuten?|Sekunden?|"
    r"Monaten?|Monat|Jahre|Jahren|Jahr|"
    r"Datensaetze[nr]?|Datensätze[nr]?|"
    r"Modelle[nr]?|Klassifikatoren?|"
    r"Faktor|fach|fache[nsr]?|malige[nsr]?|"
    r"Patient(?:innen)?|Kunden?|Nutzer(?:innen)?|"
    r"Zeilen?|Spalten?|Zertifikate?|Punkte|Punkten|"
    r"Datenquellen?|Datennodes?|Zeichen|Sensoren|"
    r"Routen|Komponenten|Entwicklern?|Releases?|"
    r"Euro|Cent"
)
# Number token: allow both . and , as decimal separator, and thousands separators.
_NUMBER = r"\d+(?:[.,]\d+)?"
BOLD_METRIC_RE = re.compile(
    r"("
    # ranges like "0.79 to 0.88" or "0,79 auf 0,88" or "44 Prozent auf 16,7 Prozent"
    + _NUMBER + r"(?:\s*(?:" + _UNIT_WORDS + r"))?\s+(?:to|auf|bis)\s+"
    + _NUMBER + r"(?:\s*(?:" + _UNIT_WORDS + r"))?"
    r"|"
    # single number with a unit word: "80 Prozent", "128 tausend", "6 Monaten"
    + _NUMBER + r"\+?\s*(?:" + _UNIT_WORDS + r")"
    r"|"
    # bare decimals with EN or DE separator: 4.4, 0,88, 1,9
    + r"\d+[.,]\d+"
    r")",
    re.IGNORECASE,
)


def wrap_bold_html(text):
    """Wrap concrete metrics in <strong> tags. Assumes text is already escaped."""
    return BOLD_METRIC_RE.sub(r"<strong>\1</strong>", text)


def bullet_runs(text):
    """Split a bullet into a list of (text, bold) tuples for DOCX rendering."""
    parts = []
    last = 0
    for m in BOLD_METRIC_RE.finditer(text):
        if m.start() > last:
            parts.append((text[last:m.start()], False))
        parts.append((m.group(1), True))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False))
    return parts if parts else [(text, False)]


# --- Shared CSS -----------------------------------------------------------
# ATS-clean template baseline, 11 August 2026. Single column entries, no
# tracked-caps CV strip, Skills section, PD/Languages as one line per field.
# Page margins moved from .page div padding to @page rule so every page (not
# just page 1) inherits the same 20mm top / 15mm bottom / 18mm side margin.
CSS = f"""
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
html {{ -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
body {{ font-family: 'Inter', 'Helvetica Neue', Arial, sans-serif; background: #ffffff; color: {BODY_HEX}; }}
.page {{ background: #ffffff; }}

.cv-tag {{ font-size: 10pt; color: {RUST_HEX}; letter-spacing: 6px; font-weight: 600; margin-bottom: 2mm; }}
h1.name {{ font-size: 22pt; font-weight: 700; color: {NAVY_HEX}; letter-spacing: 0; margin-bottom: 1.5mm; }}
.role-tag {{ font-size: 11pt; color: {NAVY_GREY_HEX}; margin-bottom: 3mm; font-weight: 500; }}
.header-rule {{ border-top: 1px solid {NAVY_HEX}; margin-bottom: 4mm; }}

.section {{ margin-bottom: 3mm; page-break-inside: avoid; break-inside: avoid; }}
.section.section-long {{ page-break-inside: auto; break-inside: auto; }}
.section h2 {{ font-size: 11.5pt; color: {NAVY_HEX}; font-weight: 700; margin-bottom: 1.5mm; text-transform: uppercase; padding-bottom: 0.5mm; border-bottom: 1px solid {RULE_HEX}; page-break-after: avoid; break-after: avoid; }}
.section h2 + .entry, .section h2 + p, .section h2 + .pd-line, .section h2 + .lang-line {{ page-break-before: avoid; break-before: avoid; }}

.pd-line {{ font-size: 10.5pt; margin-bottom: 0.4mm; color: {BODY_HEX}; }}
.pd-line .label {{ color: {NAVY_HEX}; font-weight: 700; }}

.profile-text {{ font-size: 10.5pt; line-height: 1.35; text-align: justify; color: {BODY_HEX}; }}

.skills-line {{ font-size: 10.5pt; line-height: 1.4; color: {BODY_HEX}; }}

.entry {{ margin-bottom: 3mm; page-break-inside: avoid; break-inside: avoid; }}
.entry:last-child {{ margin-bottom: 0; }}
.entry .title-line {{ font-size: 11pt; font-weight: 700; color: {NAVY_HEX}; margin-bottom: 0.3mm; }}
.entry .sub-line {{ font-size: 10pt; color: {NAVY_GREY_HEX}; margin-bottom: 1mm; font-style: italic; }}
.entry ul {{ list-style: none; padding: 0; margin-top: 0.5mm; }}
.entry ul li {{ font-size: 10.5pt; line-height: 1.3; text-align: justify; color: {BODY_HEX}; padding-left: 4mm; position: relative; margin-bottom: 0.5mm; }}
.entry ul li::before {{ content: "•"; position: absolute; left: 0; top: 0; color: {RUST_HEX}; font-weight: 700; }}
.entry .tech {{ font-size: 9.5pt; color: {NAVY_GREY_HEX}; font-style: italic; margin-top: 1mm; }}
.entry .tech .tech-label {{ font-weight: 700; font-style: normal; color: {RUST_HEX}; margin-right: 1mm; }}
.entry strong {{ color: {NAVY_HEX}; font-weight: 700; }}

.simple-list {{ list-style: none; padding: 0; }}
.simple-list li {{ font-size: 10.5pt; line-height: 1.3; text-align: justify; color: {BODY_HEX}; padding-left: 4mm; position: relative; margin-bottom: 0.4mm; }}
.simple-list li::before {{ content: "•"; position: absolute; left: 0; top: 0; color: {RUST_HEX}; font-weight: 700; }}
.simple-list strong {{ color: {NAVY_HEX}; font-weight: 700; }}

.lang-line {{ font-size: 10.5pt; margin-bottom: 0.4mm; color: {BODY_HEX}; }}
.lang-line .label {{ color: {NAVY_HEX}; font-weight: 700; }}

@page {{ size: A4; margin: 20mm 18mm 15mm 18mm; }}
"""

CL_CSS = f"""
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
html {{ -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
body {{ font-family: 'Inter', 'Helvetica Neue', Arial, sans-serif; background: #ffffff; color: {BODY_HEX}; }}
.page {{ width: 210mm; min-height: 297mm; margin: 0 auto; background: #ffffff; padding: 18mm 22mm; }}
h1 {{ font-size: 22pt; color: {NAVY_HEX}; font-weight: 700; margin-bottom: 4mm; }}
.contact {{ font-size: 10pt; color: {BODY_HEX}; margin-bottom: 8mm; }}
.date {{ font-size: 10.5pt; color: {BODY_HEX}; margin-bottom: 6mm; }}
.company {{ font-size: 11.5pt; font-weight: 700; color: {NAVY_HEX}; margin-bottom: 1mm; }}
.hiring {{ font-size: 10.5pt; color: {BODY_HEX}; margin-bottom: 6mm; }}
.subject {{ font-size: 11pt; font-weight: 700; color: {NAVY_HEX}; margin-bottom: 6mm; }}
p {{ font-size: 10.5pt; line-height: 1.5; text-align: justify; color: {BODY_HEX}; margin-bottom: 4mm; }}
.close {{ font-size: 10.5pt; color: {BODY_HEX}; margin-top: 4mm; }}
.sig {{ font-size: 11pt; font-weight: 700; color: {NAVY_HEX}; margin-top: 4mm; }}
@page {{ size: A4; margin: 0; }}
"""

# Tight variant applied by the CL auto-tightening loop when the cover letter
# renders over 1 A4 page. Shrinks page padding, line height, and paragraph
# margins so a busy cover letter still fits on one page before any content is
# dropped. Font floor of 10.5pt is preserved per the standing rule.
CL_CSS_TIGHT = f"""
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
html {{ -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
body {{ font-family: 'Inter', 'Helvetica Neue', Arial, sans-serif; background: #ffffff; color: {BODY_HEX}; }}
.page {{ width: 210mm; min-height: 297mm; margin: 0 auto; background: #ffffff; padding: 14mm 20mm; }}
h1 {{ font-size: 20pt; color: {NAVY_HEX}; font-weight: 700; margin-bottom: 3mm; }}
.contact {{ font-size: 10pt; color: {BODY_HEX}; margin-bottom: 5mm; }}
.date {{ font-size: 10.5pt; color: {BODY_HEX}; margin-bottom: 4mm; }}
.company {{ font-size: 11.5pt; font-weight: 700; color: {NAVY_HEX}; margin-bottom: 1mm; }}
.hiring {{ font-size: 10.5pt; color: {BODY_HEX}; margin-bottom: 4mm; }}
.subject {{ font-size: 11pt; font-weight: 700; color: {NAVY_HEX}; margin-bottom: 4mm; }}
p {{ font-size: 10.5pt; line-height: 1.35; text-align: justify; color: {BODY_HEX}; margin-bottom: 2.5mm; }}
.close {{ font-size: 10.5pt; color: {BODY_HEX}; margin-top: 3mm; }}
.sig {{ font-size: 11pt; font-weight: 700; color: {NAVY_HEX}; margin-top: 2mm; }}
@page {{ size: A4; margin: 0; }}
"""


def _entry_html(dates, place, title, sub, bullets, tech_stack, showcase=None):
    """ATS-clean single column entry.

    Renders title bold on one line, then italic sub line with subtitle plus
    location plus date range joined by a middle dot separator, then the
    bullets and Technologies line. No left/right grid. Naive ATS extractors
    read the title line first, then the sub line, then the bullets, in
    document order. No column pairing failure modes.
    """
    sub_parts = []
    if sub:
        sub_parts.append(sub)
    if place:
        sub_parts.append(place)
    if dates:
        sub_parts.append(dates)
    sub_line = " · ".join(sub_parts)
    sub_html = f'<div class="sub-line">{escape(sub_line)}</div>' if sub_line else ""

    bullet_html = "".join(f"<li>{wrap_bold_html(escape(b))}</li>" for b in bullets)
    ul_html = f"<ul>{bullet_html}</ul>" if bullets else ""

    tech_html = ""
    if tech_stack:
        tech_html = (
            f'<div class="tech"><span class="tech-label">Technologies:</span>'
            f'{escape(", ".join(tech_stack))}</div>'
        )
    live_html = ""
    if showcase:
        live_html = (
            f'<div class="tech"><span class="tech-label">Live:</span>'
            f'<a href="https://{escape(showcase)}" style="color: inherit; text-decoration: none;">{escape(showcase)}</a></div>'
        )

    return (
        f'<div class="entry">'
        f'<div class="title-line">{escape(title)}</div>'
        f'{sub_html}{ul_html}{tech_html}{live_html}'
        f'</div>'
    )


def _entry_head_only_html(dates, place, title, sub):
    """ATS-clean head-only entry, used for Education. Same shape as _entry_html
    but no bullets and no Technologies line."""
    return _entry_html(dates, place, title, sub, [], [])


DEFAULT_SKILLS = [
    "Python", "SQL", "PySpark", "BigQuery", "Databricks", "Delta Lake",
    "Power BI", "Tableau", "Looker Studio", "LangChain", "LangGraph",
    "Ollama", "CatBoost", "LightGBM", "XGBoost", "Prophet", "MICE",
    "scikit-learn", "PyTorch", "dbt", "Apache Airflow", "GCP", "AWS",
    "Docker", "Git", "React", "Playwright",
]


def _skills_line(cfg):
    """Return the comma separated Skills line for the ATS Skills/Fähigkeiten
    section. Uses cfg['skills'] when the config overrides, otherwise the
    curated DEFAULT_SKILLS list. Section heading changes between EN and DE
    tracks; the skill tokens themselves stay in English as they are used."""
    skills = cfg.get("skills") or DEFAULT_SKILLS
    return ", ".join(skills)


def html_cv(cfg):
    is_de = cfg.get("lang") == "de"
    hdr = HDR_DE if is_de else HDR_EN
    pd_fields = PD_FIELDS_DE if is_de else PD_FIELDS_EN
    skills_heading = "Fähigkeiten" if is_de else "Skills"

    pd_html = "".join(
        f'<div class="pd-line"><span class="label">{escape(lbl)}:</span> {escape(val)}</div>'
        for lbl, val in pd_fields
    )

    # Experience, three entries per the 2 August 2026 rule, reverse chronological.
    # 1. eRay GmbH Data Scientist. 2. SS Engineers Junior Associate Software
    # Developer. 3. SS Engineers Front End Developer Intern. SS Engineers bullets
    # come from role_configs.py so master-projects.md remains the source of truth.
    from role_configs import (
        SATENDRA_FT_BULLETS_EN, SATENDRA_FT_BULLETS_DE,
        SATENDRA_INTERN_BULLETS_EN, SATENDRA_INTERN_BULLETS_DE,
    )
    exp_dates = "Oct 2025 to Mar 2026" if not is_de else "Okt 2025 bis Mrz 2026"
    exp_place = "Heidelberg"
    ft_dates = "Aug 2023 to Aug 2024" if not is_de else "Aug 2023 bis Aug 2024"
    intern_dates = "Feb 2023 to July 2023" if not is_de else "Feb 2023 bis Juli 2023"
    ft_title = "Junior Associate Software Developer"
    intern_title = "Front End Developer Intern"
    ft_all = SATENDRA_FT_BULLETS_DE if is_de else SATENDRA_FT_BULLETS_EN
    intern_all = SATENDRA_INTERN_BULLETS_DE if is_de else SATENDRA_INTERN_BULLETS_EN
    # 4 August 2026 CV length calibration defaults: SS FT capped at 2 bullets,
    # SS Intern capped at 1 bullet. Overridable via cfg for tighter overflow.
    ft_bullets = ft_all[:cfg.get("ss_ft_max_bullets", 2)]
    intern_bullets = intern_all[:cfg.get("ss_intern_max_bullets", 1)]
    exp_html = (
        _entry_html(
            exp_dates, exp_place,
            "eRay GmbH", "Data Scientist",
            cfg["experience_bullets"],
            ["Python", "CatBoost", "Prophet", "scikit learn", "MICE"],
        )
        + _entry_html(
            ft_dates, "India",
            "SS Engineers and Contractors", ft_title,
            ft_bullets,
            ["React", "module federation", "Playwright", "AngularJS", "HTML5", "CSS3"],
        )
        + _entry_html(
            intern_dates, "India",
            "SS Engineers and Contractors", intern_title,
            intern_bullets,
            ["React", "HTML5", "CSS3", "Git", "code review workflow"],
        )
    )

    # Education
    edu_dates_1 = "Apr 2025 to Present" if not is_de else "Apr 2025 bis heute"
    edu_dates_2 = "2019 to 2023"
    edu_html = (
        _entry_head_only_html(
            edu_dates_1, "Heidelberg",
            "M.Sc. Data Science and Analytics",
            "SRH University of Applied Sciences Heidelberg, GPA 1.9",
        )
        + _entry_head_only_html(
            edu_dates_2, "Greater Noida, India",
            "Bachelor of Technology in Computer Science",
            "GL Bajaj Institute of Technology and Management, CGPA 7.3 of 10",
        )
    )

    # Personal Projects
    proj_html = ""
    pp_max = cfg.get("pp_bullets_per_entry", 3)
    for p in cfg["projects"]:
        proj_html += _entry_html(
            p.get("date_label", "2025"),
            p.get("kind_label", ""),
            p["title"],
            "",
            p["bullets"][:pp_max],
            p["stack"],
            p.get("showcase"),
        )

    # Research
    research_title = (
        "Bachelor Thesis, Diabetes Prediction Using Machine Learning"
        if not is_de else "Bachelorarbeit: Diabetesvorhersage mit Machine Learning"
    )
    research_place = "Greater Noida, India"
    research_html = _entry_html(
        "2022 to 2023", research_place,
        research_title,
        "GL Bajaj Institute of Technology and Management, IEEE style paper",
        cfg["research_bullets"],
        ["Python", "scikit learn", "Pandas", "Seaborn"],
    )

    # Certifications, achievements
    certs_html = (
        '<ul class="simple-list">'
        + "".join(f"<li>{escape(c)}</li>" for c in cfg["certifications"])
        + "</ul>"
    )
    ach_html = (
        '<ul class="simple-list">'
        + "".join(f"<li>{wrap_bold_html(escape(a))}</li>" for a in cfg["achievements"])
        + "</ul>"
    )

    # Languages, one row per language so each is scannable on its own line
    if is_de:
        lang_fields = [
            ("Englisch", "Fließend, schriftlich und mündlich"),
            ("Deutsch", "B1 laufend Richtung B2"),
            ("Hindi", "Muttersprache"),
        ]
    else:
        lang_fields = [
            ("English", "Fluent, written and spoken"),
            ("German", "B1 in progress toward B2"),
            ("Hindi", "Native"),
        ]
    lang_html = "".join(
        f'<div class="lang-line"><span class="label">{escape(lbl)}:</span> {escape(val)}</div>'
        for lbl, val in lang_fields
    )

    # ATS-clean section ordering per 11 August 2026 rule:
    # PD -> Profile -> Skills -> Experience -> Education -> Projects ->
    # Research -> Certs -> Achievements -> Languages.
    skills_html = f'<p class="skills-line">{escape(_skills_line(cfg))}</p>'

    return f"""<!DOCTYPE html><html lang="{'de' if is_de else 'en'}"><head><meta charset="UTF-8"><title>Rahul Rawat, {escape(cfg['role_strip'])}</title><style>{CSS}</style></head><body><div class="page">
<!-- CURRICULUM VITAE strip retired 11 August 2026 for ATS parseability -->
<h1 class="name">Rahul Rawat</h1>
<div class="role-tag">{escape(cfg['role_strip'])}</div>
<div class="header-rule"></div>
<section class="section"><h2>{escape(hdr['pd'])}</h2>{pd_html}</section>
<section class="section"><h2>{escape(hdr['profile'])}</h2><p class="profile-text">{wrap_bold_html(escape(cfg['profile']))}</p></section>
<section class="section"><h2>{escape(skills_heading)}</h2>{skills_html}</section>
<section class="section section-long"><h2>{escape(hdr['experience'])}</h2>{exp_html}</section>
<section class="section"><h2>{escape(hdr['education'])}</h2>{edu_html}</section>
<section class="section section-long"><h2>{escape(hdr['projects'])}</h2>{proj_html}</section>
<section class="section"><h2>{escape(hdr['research'])}</h2>{research_html}</section>
<section class="section"><h2>{escape(hdr['certifications'])}</h2>{certs_html}</section>
<section class="section"><h2>{escape(hdr['achievements'])}</h2>{ach_html}</section>
<section class="section"><h2>{escape(hdr['languages'])}</h2>{lang_html}</section>
</div></body></html>"""


def html_cover_letter(cfg, tight=False):
    is_de = cfg.get("lang") == "de"
    if is_de:
        contact_line = "rahulrawat2r@gmail.com · 015563603340 · linkedin.com/in/rahulrawat2r · github.com/rahulrawat20022002 · Mannheim, Deutschland"
        hiring = "Personalabteilung"
        subject = f"Bewerbung: {escape(cfg['cl_subject'])}"
        greeting = "Sehr geehrte Damen und Herren,"
        close_line = "Mit freundlichen Grüßen,"
    else:
        contact_line = "rahulrawat2r@gmail.com · 015563603340 · linkedin.com/in/rahulrawat2r · github.com/rahulrawat20022002 · Mannheim, Germany"
        hiring = "Hiring Team"
        subject = f"Application: {escape(cfg['cl_subject'])}"
        greeting = "Dear Hiring Team,"
        close_line = "Kind regards,"

    paras = "".join(f"<p>{escape(p)}</p>" for p in cfg["cl_paragraphs"])
    css = CL_CSS_TIGHT if tight else CL_CSS
    return f"""<!DOCTYPE html><html lang="{'de' if is_de else 'en'}"><head><meta charset="UTF-8"><title>Cover Letter Rahul Rawat</title><style>{css}</style></head><body><div class="page"><h1>Rahul Rawat</h1><div class="contact">{contact_line}</div><div class="date">{escape(cfg['cl_date'])}</div><div class="company">{escape(cfg['company'])}</div><div class="hiring">{hiring}</div><div class="subject">{subject}</div><p>{greeting}</p>{paras}<p class="close">{close_line}</p><p class="sig">Rahul Rawat</p></div></body></html>"""


# --- DOCX rendering ------------------------------------------------------
def _set_para_border_bottom(paragraph, color="B7C0CE"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _zero_cell_margins(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for edge in ('top', 'left', 'bottom', 'right'):
        m = OxmlElement(f'w:{edge}')
        m.set(qn('w:w'), '0')
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)


def _add_section_heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(" ".join(text.upper()))
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    _set_para_border_bottom(p)
    return p


def _mark_row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    trPr.append(cantSplit)


def _add_pd_line(doc, label, value):
    """ATS-clean PD entry as a plain paragraph: 'Label: Value'. Replaces the
    old two-column table so naive ATS parsers do not read labels first and
    values second."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    lr = p.add_run(f"{label}: ")
    lr.bold = True
    lr.font.size = Pt(10.5)
    lr.font.color.rgb = NAVY
    vr = p.add_run(value)
    vr.font.size = Pt(10.5)
    vr.font.color.rgb = BODY


def _add_bulleted_runs(paragraph, text):
    # First a rust bullet mark
    br = paragraph.add_run("• ")
    br.bold = True
    br.font.size = Pt(10.5)
    br.font.color.rgb = RUST
    for chunk, is_bold in bullet_runs(text):
        rr = paragraph.add_run(chunk)
        rr.font.size = Pt(10.5)
        rr.bold = is_bold
        rr.font.color.rgb = NAVY if is_bold else BODY


def _add_entry(doc, dates, place, title, sub, bullets, tech_stack, showcase=None):
    """ATS-clean single-column DOCX entry.

    Title line bold on its own paragraph, then italic sub line joining
    subtitle + location + date range with a middle dot separator, then
    bullets, then Technologies line. No two-column tables to prevent ATS
    parsers from scrambling label-value pairs when they read column-by-column.
    """
    # Title line
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(2)
    tp.paragraph_format.space_after = Pt(0)
    tp.paragraph_format.line_spacing = 1.15
    tp.paragraph_format.keep_with_next = True
    tr = tp.add_run(title)
    tr.bold = True
    tr.font.size = Pt(11)
    tr.font.color.rgb = NAVY

    # Sub line: "<sub> · <place> · <dates>"
    sub_parts = [p for p in (sub, place, dates) if p]
    if sub_parts:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(0)
        sp.paragraph_format.line_spacing = 1.15
        sp.paragraph_format.keep_with_next = True
        sr = sp.add_run(" · ".join(sub_parts))
        sr.italic = True
        sr.font.size = Pt(10)
        sr.font.color.rgb = NAVY_GREY

    for b in bullets:
        bp = doc.add_paragraph()
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after = Pt(0)
        bp.paragraph_format.line_spacing = 1.15
        bp.paragraph_format.left_indent = Cm(0.5)
        bp.paragraph_format.first_line_indent = Cm(-0.3)
        bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_bulleted_runs(bp, b)

    if tech_stack:
        tp2 = doc.add_paragraph()
        tp2.paragraph_format.space_before = Pt(1)
        tp2.paragraph_format.space_after = Pt(0)
        lbl = tp2.add_run("Technologies: ")
        lbl.bold = True
        lbl.font.size = Pt(9.5)
        lbl.font.color.rgb = RUST
        val = tp2.add_run(", ".join(tech_stack))
        val.italic = True
        val.font.size = Pt(9.5)
        val.font.color.rgb = NAVY_GREY

    if showcase:
        lp2 = doc.add_paragraph()
        lp2.paragraph_format.space_before = Pt(0.5)
        lp2.paragraph_format.space_after = Pt(0)
        lbl2 = lp2.add_run("Live: ")
        lbl2.bold = True
        lbl2.font.size = Pt(9.5)
        lbl2.font.color.rgb = RUST
        val2 = lp2.add_run(showcase)
        val2.italic = True
        val2.font.size = Pt(9.5)
        val2.font.color.rgb = NAVY_GREY


def docx_cv(cfg, path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.6)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)

    is_de = cfg.get("lang") == "de"
    hdr = HDR_DE if is_de else HDR_EN
    pd_fields = PD_FIELDS_DE if is_de else PD_FIELDS_EN

    # Header, compact. CURRICULUM VITAE strip retired 11 August 2026 for ATS
    # parseability. Naive PDF text extractors split the spaced-caps tag into
    # one letter per line, which many ATS parsers count as noise before the
    # name.
    name_p = doc.add_paragraph()
    name_p.paragraph_format.space_before = Pt(0)
    name_p.paragraph_format.space_after = Pt(0)
    name_p.paragraph_format.line_spacing = 1.0
    name_r = name_p.add_run("Rahul Rawat")
    name_r.bold = True
    name_r.font.size = Pt(20)
    name_r.font.color.rgb = NAVY

    tagline_p = doc.add_paragraph()
    tagline_p.paragraph_format.space_before = Pt(0)
    tagline_p.paragraph_format.space_after = Pt(1)
    tagline_p.paragraph_format.line_spacing = 1.0
    tagline_r = tagline_p.add_run(cfg["role_strip"])
    tagline_r.font.size = Pt(10.5)
    tagline_r.font.color.rgb = NAVY_GREY
    _set_para_border_bottom(tagline_p, color=NAVY_HEX.replace("#", ""))

    # Personal Details, one line per field for ATS parseability
    _add_section_heading(doc, hdr["pd"])
    for lbl, val in pd_fields:
        _add_pd_line(doc, lbl, val)

    # Profile
    _add_section_heading(doc, hdr["profile"])
    prof_p = doc.add_paragraph()
    prof_p.paragraph_format.space_before = Pt(0)
    prof_p.paragraph_format.space_after = Pt(0)
    prof_p.paragraph_format.line_spacing = 1.2
    prof_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for chunk, is_bold in bullet_runs(cfg["profile"]):
        r = prof_p.add_run(chunk)
        r.font.size = Pt(10.5)
        r.bold = is_bold
        r.font.color.rgb = NAVY if is_bold else BODY

    # Skills / Fähigkeiten, one comma separated line for ATS keyword coverage
    skills_heading = "Fähigkeiten" if is_de else "Skills"
    _add_section_heading(doc, skills_heading)
    sk_p = doc.add_paragraph()
    sk_p.paragraph_format.space_before = Pt(0)
    sk_p.paragraph_format.space_after = Pt(0)
    sk_p.paragraph_format.line_spacing = 1.25
    sk_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sk_r = sk_p.add_run(_skills_line(cfg))
    sk_r.font.size = Pt(10.5)
    sk_r.font.color.rgb = BODY

    # Professional Experience, three entries per the 2 August 2026 rule
    from role_configs import (
        SATENDRA_FT_BULLETS_EN, SATENDRA_FT_BULLETS_DE,
        SATENDRA_INTERN_BULLETS_EN, SATENDRA_INTERN_BULLETS_DE,
    )
    _add_section_heading(doc, hdr["experience"])
    exp_dates = "Oct 2025 to Mar 2026" if not is_de else "Okt 2025 bis Mrz 2026"
    ft_dates = "Aug 2023 to Aug 2024" if not is_de else "Aug 2023 bis Aug 2024"
    intern_dates = "Feb 2023 to July 2023" if not is_de else "Feb 2023 bis Juli 2023"
    ft_all = SATENDRA_FT_BULLETS_DE if is_de else SATENDRA_FT_BULLETS_EN
    intern_all = SATENDRA_INTERN_BULLETS_DE if is_de else SATENDRA_INTERN_BULLETS_EN
    # 4 August 2026 CV length calibration defaults: SS FT capped at 2 bullets,
    # SS Intern capped at 1 bullet. Overridable via cfg for tighter overflow.
    ft_bullets = ft_all[:cfg.get("ss_ft_max_bullets", 2)]
    intern_bullets = intern_all[:cfg.get("ss_intern_max_bullets", 1)]
    _add_entry(
        doc, exp_dates, "Heidelberg",
        "eRay GmbH", "Data Scientist",
        cfg["experience_bullets"],
        ["Python", "CatBoost", "Prophet", "scikit learn", "MICE"],
    )
    _add_entry(
        doc, ft_dates, "India",
        "SS Engineers and Contractors", "Junior Associate Software Developer",
        ft_bullets,
        ["React", "module federation", "Playwright", "AngularJS", "HTML5", "CSS3"],
    )
    _add_entry(
        doc, intern_dates, "India",
        "SS Engineers and Contractors", "Front End Developer Intern",
        intern_bullets,
        ["React", "HTML5", "CSS3", "Git", "code review workflow"],
    )

    # Education, placed right after Experience per ATS-clean section order
    _add_section_heading(doc, hdr["education"])
    edu_dates_1 = "Apr 2025 to Present" if not is_de else "Apr 2025 bis heute"
    _add_entry(
        doc, edu_dates_1, "Heidelberg",
        "M.Sc. Data Science and Analytics",
        "SRH University of Applied Sciences Heidelberg, GPA 1.9",
        [], [],
    )
    _add_entry(
        doc, "2019 to 2023", "Greater Noida, India",
        "Bachelor of Technology in Computer Science",
        "GL Bajaj Institute of Technology and Management, CGPA 7.3 of 10",
        [], [],
    )

    # Personal Projects
    _add_section_heading(doc, hdr["projects"])
    pp_max = cfg.get("pp_bullets_per_entry", 3)
    for p in cfg["projects"]:
        _add_entry(
            doc,
            p.get("date_label", "2025"),
            p.get("kind_label", ""),
            p["title"], "",
            p["bullets"][:pp_max],
            p["stack"],
            p.get("showcase"),
        )

    # Research
    _add_section_heading(doc, hdr["research"])
    research_title = (
        "Bachelor Thesis, Diabetes Prediction Using Machine Learning"
        if not is_de else "Bachelorarbeit: Diabetesvorhersage mit Machine Learning"
    )
    _add_entry(
        doc, "2022 to 2023", "Greater Noida, India",
        research_title,
        "GL Bajaj Institute of Technology and Management, IEEE style paper",
        cfg["research_bullets"],
        ["Python", "scikit learn", "Pandas", "Seaborn"],
    )

    # Certifications
    _add_section_heading(doc, hdr["certifications"])
    for c in cfg["certifications"][:3]:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(0)
        cp.paragraph_format.line_spacing = 1.05
        cp.paragraph_format.left_indent = Cm(0.3)
        _add_bulleted_runs(cp, c)

    # Achievements
    _add_section_heading(doc, hdr["achievements"])
    for a in cfg["achievements"]:
        ap = doc.add_paragraph()
        ap.paragraph_format.space_before = Pt(0)
        ap.paragraph_format.space_after = Pt(0)
        ap.paragraph_format.line_spacing = 1.1
        ap.paragraph_format.left_indent = Cm(0.3)
        _add_bulleted_runs(ap, a)

    # Languages, one row per language so each is scannable on its own line
    _add_section_heading(doc, hdr["languages"])
    if is_de:
        lang_fields = [
            ("Englisch", "Fließend, schriftlich und mündlich"),
            ("Deutsch", "B1 laufend Richtung B2"),
            ("Hindi", "Muttersprache"),
        ]
    else:
        lang_fields = [
            ("English", "Fluent, written and spoken"),
            ("German", "B1 in progress toward B2"),
            ("Hindi", "Native"),
        ]
    for lbl, val in lang_fields:
        _add_pd_line(doc, lbl, val)

    # Signature block retired 4 August 2026, CVs no longer carry Mannheim,
    # date, and name at the bottom. Cover letters keep their own closing.

    doc.save(str(path))


def docx_cover_letter(cfg, path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

    is_de = cfg.get("lang") == "de"

    # Header name
    p = doc.add_paragraph()
    r = p.add_run("Rahul Rawat")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = NAVY

    contact_line = (
        "rahulrawat2r@gmail.com  ·  015563603340  ·  linkedin.com/in/rahulrawat2r  ·  github.com/rahulrawat20022002  ·  Mannheim, "
        + ("Deutschland" if is_de else "Germany")
    )
    p = doc.add_paragraph()
    r = p.add_run(contact_line)
    r.font.size = Pt(10)
    r.font.color.rgb = BODY

    p = doc.add_paragraph()
    r = p.add_run(cfg["cl_date"])
    r.font.size = Pt(10.5)
    r.font.color.rgb = BODY

    p = doc.add_paragraph()
    r = p.add_run(cfg["company"])
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    r = p.add_run("Personalabteilung" if is_de else "Hiring Team")
    r.font.size = Pt(10.5)
    r.font.color.rgb = BODY

    p = doc.add_paragraph()
    r = p.add_run(("Bewerbung: " if is_de else "Application: ") + cfg["cl_subject"])
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    r = p.add_run("Sehr geehrte Damen und Herren," if is_de else "Dear Hiring Team,")
    r.font.size = Pt(10.5)
    r.font.color.rgb = BODY

    for para in cfg["cl_paragraphs"]:
        pp = doc.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pp.paragraph_format.line_spacing = 1.35
        rr = pp.add_run(para)
        rr.font.size = Pt(10.5)
        rr.font.color.rgb = BODY

    p = doc.add_paragraph()
    r = p.add_run("Mit freundlichen Grüßen," if is_de else "Kind regards,")
    r.font.size = Pt(10.5)
    r.font.color.rgb = BODY

    p = doc.add_paragraph()
    r = p.add_run("Rahul Rawat")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = NAVY

    doc.save(str(path))


def build_role(cfg):
    folder = DRAFTS / cfg["folder"]
    folder.mkdir(parents=True, exist_ok=True)

    cv_html_path = folder / "CV_Rahul_Rawat.html"
    cv_pdf_path = folder / "CV_Rahul_Rawat.pdf"
    cv_docx_path = folder / "CV_Rahul_Rawat.docx"
    cl_html_path = folder / "CoverLetter_Rahul_Rawat.html"
    cl_pdf_path = folder / "CoverLetter_Rahul_Rawat.pdf"
    cl_docx_path = folder / "CoverLetter_Rahul_Rawat.docx"
    cl_md_path = folder / "CoverLetter_Rahul_Rawat.md"
    cv_md_path = folder / "CV_Rahul_Rawat.md"

    # 4 August 2026 hard 3 page cap. Overflow ladder: drop trailing Personal
    # Projects entries first (configs are ordered by relevance, last is least
    # relevant), then trim SS Engineers bullets (FT down from 4 to 2, then
    # Intern down from 2 to 1). Per the 2 August 2026 rule, all three
    # Experience entries stay; only bullet counts inside them shrink.
    import copy, pypdf
    # 4 August 2026 CV length calibration defaults: Personal Projects capped
    # at 2 entries max. Cap here before the ladder so all render paths use it.
    if len(cfg.get("projects", [])) > 2:
        cfg["projects"] = cfg["projects"][:2]
    working_cfg = copy.deepcopy(cfg)
    original_project_count = len(working_cfg.get("projects", []))
    cv_html = None
    pages = None
    # Ladder rungs, applied in order until the CV fits 3 pages. Starts from
    # the defaults (2 projects, SS FT 2 bullets, SS intern 1 bullet) and only
    # tightens further from there.
    original_cert_count = len(cfg.get("certifications", []))
    original_research_count = len(cfg.get("research_bullets", []))
    ladder = []
    # Start from defaults: up to 2 projects at 3 bullets each, SS FT 2, SS
    # intern 1, all certs, all research bullets. Ladder rungs tighten from here.
    for pc in range(original_project_count, 0, -1):
        ladder.append({"projects_kept": pc, "pp_bullets": 3, "ss_ft": 2, "ss_intern": 1, "certs_kept": original_cert_count, "research_kept": original_research_count})
    # Trim PP bullets from 3 to 2 while keeping projects_kept at original
    for pc in range(original_project_count, 0, -1):
        ladder.append({"projects_kept": pc, "pp_bullets": 2, "ss_ft": 2, "ss_intern": 1, "certs_kept": original_cert_count, "research_kept": original_research_count})
    # Trim SS FT further at 1 project and reduced PP bullets
    ladder.append({"projects_kept": 1, "pp_bullets": 2, "ss_ft": 1, "ss_intern": 1, "certs_kept": original_cert_count, "research_kept": original_research_count})
    # Then drop certificates one at a time
    for ck in range(original_cert_count - 1, 0, -1):
        ladder.append({"projects_kept": 1, "pp_bullets": 2, "ss_ft": 1, "ss_intern": 1, "certs_kept": ck, "research_kept": original_research_count})
    # Then trim research bullets
    for rk in range(original_research_count - 1, 1, -1):
        ladder.append({"projects_kept": 1, "pp_bullets": 2, "ss_ft": 1, "ss_intern": 1, "certs_kept": max(1, original_cert_count - 2), "research_kept": rk})
    tried = []
    try:
        for step in ladder:
            projects = cfg.get("projects", [])[:step["projects_kept"]]
            working_cfg["projects"] = projects
            working_cfg["pp_bullets_per_entry"] = step["pp_bullets"]
            working_cfg["ss_ft_max_bullets"] = step["ss_ft"]
            working_cfg["ss_intern_max_bullets"] = step["ss_intern"]
            working_cfg["certifications"] = cfg.get("certifications", [])[:step["certs_kept"]]
            working_cfg["research_bullets"] = cfg.get("research_bullets", [])[:step["research_kept"]]
            cv_html = html_cv(working_cfg)
            weasyprint.HTML(string=cv_html).write_pdf(str(cv_pdf_path))
            pages = len(pypdf.PdfReader(str(cv_pdf_path)).pages)
            tried.append((step, pages))
            if pages <= 3:
                break
        else:
            raise RuntimeError(
                f"CV PDF for {cfg['folder']} still {pages} pages after full "
                f"overflow ladder. Steps tried: {tried}. Tighten the profile "
                f"paragraph in the config or drop a certificate for this role."
            )
        cv_html_path.write_text(cv_html, encoding="utf-8")
        first_step, first_pages = tried[0]
        last_step, last_pages = tried[-1]
        if len(tried) > 1:
            print(
                f"  auto trimmed for {cfg['folder']}: "
                f"projects {original_project_count} to {last_step['projects_kept']}, "
                f"SS FT bullets {last_step['ss_ft']}, SS intern bullets {last_step['ss_intern']}, "
                f"final pages: {last_pages}"
            )
        if last_pages < 2:
            print(f"  WARNING: {cfg['folder']} CV PDF has only {last_pages} page(s); target is 2 to 3")
        # Mutate cfg so the docx render matches the html render
        cfg["projects"] = working_cfg["projects"]
        cfg["pp_bullets_per_entry"] = working_cfg["pp_bullets_per_entry"]
        cfg["ss_ft_max_bullets"] = working_cfg["ss_ft_max_bullets"]
        cfg["ss_intern_max_bullets"] = working_cfg["ss_intern_max_bullets"]
        cfg["certifications"] = working_cfg["certifications"]
        cfg["research_bullets"] = working_cfg["research_bullets"]
    except Exception as e:
        print(f"  PDF fail for {cfg['folder']}: {e}")
        raise

    try:
        docx_cv(cfg, cv_docx_path)
    except Exception as e:
        print(f"  DOCX CV fail for {cfg['folder']}: {e}")

    try:
        docx_cover_letter(cfg, cl_docx_path)
    except Exception as e:
        print(f"  DOCX CL fail for {cfg['folder']}: {e}")

    # 4 August 2026 cover letter 1 page cap. Render the CL as HTML + PDF, and
    # if it overflows, switch to the tight CSS variant. If still over one page,
    # drop the trailing cover letter paragraph in the docx and md variants too.
    # docx and md are the primary application deliverables; the PDF and HTML
    # are decorative previews so Rah can eyeball the page layout.
    try:
        cl_working_cfg = copy.deepcopy(cfg)
        cl_pages = None
        # Rung 1, default CL CSS
        cl_html = html_cover_letter(cl_working_cfg, tight=False)
        weasyprint.HTML(string=cl_html).write_pdf(str(cl_pdf_path))
        cl_pages = len(pypdf.PdfReader(str(cl_pdf_path)).pages)
        cl_used_tight = False
        cl_trimmed = 0
        # Rung 2, tight CL CSS
        if cl_pages > 1:
            cl_used_tight = True
            cl_html = html_cover_letter(cl_working_cfg, tight=True)
            weasyprint.HTML(string=cl_html).write_pdf(str(cl_pdf_path))
            cl_pages = len(pypdf.PdfReader(str(cl_pdf_path)).pages)
        # Rung 3+, drop trailing paragraphs one at a time under tight CSS
        while cl_pages > 1 and len(cl_working_cfg.get("cl_paragraphs", [])) > 2:
            cl_working_cfg["cl_paragraphs"] = cl_working_cfg["cl_paragraphs"][:-1]
            cl_trimmed += 1
            cl_html = html_cover_letter(cl_working_cfg, tight=True)
            weasyprint.HTML(string=cl_html).write_pdf(str(cl_pdf_path))
            cl_pages = len(pypdf.PdfReader(str(cl_pdf_path)).pages)
        cl_html_path.write_text(cl_html, encoding="utf-8")
        if cl_pages > 1:
            print(
                f"  WARNING: {cfg['folder']} CoverLetter still {cl_pages} pages after tight CSS "
                f"and trimming to {len(cl_working_cfg['cl_paragraphs'])} paragraphs; "
                f"tighten paragraphs manually in the config"
            )
        elif cl_used_tight or cl_trimmed:
            print(
                f"  CL fit to 1 page for {cfg['folder']}: "
                f"tight css={cl_used_tight}, paragraphs trimmed={cl_trimmed}"
            )
        # Mirror the trimmed cl_paragraphs into cfg so docx CL and md CL match
        if cl_trimmed:
            cfg["cl_paragraphs"] = cl_working_cfg["cl_paragraphs"]
            # Rebuild docx CL with trimmed content
            docx_cover_letter(cfg, cl_docx_path)
    except Exception as e:
        print(f"  CL PDF fail for {cfg['folder']}: {e}")

    md_cv = (
        f"# Rahul Rawat\n\n## {cfg['role_strip']}\n\n{cfg['profile']}\n\n### Experience\n\n**eRay GmbH**, Data Scientist, Oct 2025 to Mar 2026\n\n"
        + "\n".join(f"* {b}" for b in cfg["experience_bullets"])
        + "\n"
    )
    cv_md_path.write_text(md_cv, encoding="utf-8")

    is_de = cfg.get("lang") == "de"
    hiring = "Personalabteilung" if is_de else "Hiring Team"
    subject_label = "Bewerbung: " if is_de else "Application: "
    greeting = "Sehr geehrte Damen und Herren," if is_de else "Dear Hiring Team,"
    close_line = "Mit freundlichen Grüßen," if is_de else "Kind regards,"
    md_cl = (
        f"# Rahul Rawat\n\nrahulrawat2r@gmail.com  ·  015563603340  ·  linkedin.com/in/rahulrawat2r  ·  github.com/rahulrawat20022002  ·  Mannheim, "
        + ("Deutschland" if is_de else "Germany")
        + f"\n\n{cfg['cl_date']}\n\n**{cfg['company']}**\n{hiring}\n\n**{subject_label}{cfg['cl_subject']}**\n\n{greeting}\n\n"
        + "\n\n".join(cfg["cl_paragraphs"])
        + f"\n\n{close_line}\n\n**Rahul Rawat**\n"
    )
    cl_md_path.write_text(md_cl, encoding="utf-8")

    print(f"  built: {cfg['folder']} ({cfg.get('lang', 'en')})")


if __name__ == "__main__":
    from role_configs import CONFIGS
    for cfg in CONFIGS:
        build_role(cfg)
